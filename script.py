import os
import shutil
import comtypes.client
import pythoncom  # Fixes "CoInitialize has not been called"
from PIL import Image
from fpdf import FPDF
from docx import Document
import re
import os 
import fitz #PyMuPDF for extracting text from PDFs
import pandas as pd
import spacy
from pdfminer.high_level import extract_text
from docx import Document
from PyPDF2 import PdfReader
import pdfplumber
from fuzzywuzzy import fuzz
from glob import glob
from datetime import datetime
import numpy as np
import streamlit as st



# Define Custom Temporary Folder for PDFs
TEMP_PDF_FOLDER = r"C:\Users\inc3061\OneDrive - Texila American University\Documents\Resumepath\Data_Flask_Task\New_Parser\Functions\processed_pdfs"
JD_UPLOAD_FOLDER = r"C:\Users\inc3061\OneDrive - Texila American University\Documents\Resumepath\Data_Flask_Task\New_Parser\Functions\JD_uploads"
os.makedirs(TEMP_PDF_FOLDER, exist_ok=True)  # Ensure folder exists
os.makedirs(JD_UPLOAD_FOLDER, exist_ok=True)

# Font Path (Ensure the font file exists)
FONT_PATH = r"C:\Users\inc3061\Downloads\dejavu-fonts-ttf-2.37\ttf\DejaVuSans.ttf"

class UnicodePDF(FPDF):
    """ Custom FPDF class to handle Unicode characters properly. """
    def header(self):
        self.set_font("Arial", "", 12)
        self.cell(0, 10, "Converted PDF", ln=True, align="C")

class FileConverter:
    def __init__(self, output_folder):
        self.output_folder = output_folder
        os.makedirs(output_folder, exist_ok=True)

    def doc_to_pdf(self, doc_path, output_pdf_path):
        """ Converts DOC/DOCX to PDF using Microsoft Word COM (Windows-only). """
        try:
            pythoncom.CoInitialize()  # Ensure COM is initialized
            word = comtypes.client.CreateObject("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            doc = word.Documents.Open(os.path.abspath(doc_path))

            doc.SaveAs(os.path.abspath(output_pdf_path), FileFormat=17)
            doc.Close(False)
            word.Quit()

            if not os.path.exists(output_pdf_path):
                raise FileNotFoundError(f"Expected PDF not found: {output_pdf_path}")

        except Exception as e:
            print(f"MS Word failed, using FPDF fallback... Error: {e}")
            self.docx_to_pdf_fallback(doc_path, output_pdf_path)

    def docx_to_pdf_fallback(self, docx_path, output_pdf_path):
        """ Fallback method using FPDF if Word COM fails. """
        try:
            doc = Document(docx_path)
            pdf = UnicodePDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()

            # Add font safely
            if os.path.exists(FONT_PATH):
                pdf.add_font("DejaVu", "", FONT_PATH, uni=True)
                pdf.set_font("DejaVu", "", 12)
            else:
                pdf.set_font("Arial", "", 12)  # Fallback to Arial

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    pdf.multi_cell(190, 10, text)

            pdf.output(output_pdf_path, "F")

        except Exception as e:
            print(f"FPDF conversion failed: {e}")

    def convert_to_pdf(self, input_file, file_ext):
        """ Converts various files (PDF, Image, DOC, DOCX) to PDF. """
        output_pdf_path = os.path.join(self.output_folder, os.path.splitext(os.path.basename(input_file))[0] + ".pdf")

        try:
            # If input file is already a PDF in the same folder, do nothing
            if file_ext == "pdf":
                if os.path.abspath(input_file) == os.path.abspath(output_pdf_path):
                    return output_pdf_path  # Avoid copying to itself
                shutil.copy(input_file, output_pdf_path)

            elif file_ext in ["jpg", "jpeg", "png"]:
                img = Image.open(input_file)
                img.convert("RGB").save(output_pdf_path, "PDF")

            elif file_ext in ["doc", "docx"]:
                self.doc_to_pdf(input_file, output_pdf_path)

            else:
                return f"Unsupported file format: {file_ext}"

            return output_pdf_path if os.path.exists(output_pdf_path) else f"Conversion failed: {output_pdf_path}"

        except Exception as e:
            return f"Error during conversion: {str(e)}"

#Streamlit App
st.title("Resume Parser")

# File Uploader for PDF Conversion
st.header("File for PDF Conversion")
uploaded_file = st.file_uploader("Upload a file", type=["pdf", "jpg", "jpeg", "png", "doc", "docx"])
if uploaded_file:
    file_ext = uploaded_file.name.split(".")[-1].lower()
    temp_path = os.path.join(TEMP_PDF_FOLDER, uploaded_file.name)

    # Save uploaded file to the custom temp folder
    with open(temp_path, "wb") as f:
        f.write(uploaded_file.getbuffer())

    converter = FileConverter(TEMP_PDF_FOLDER)
    converted_pdf = converter.convert_to_pdf(temp_path, file_ext)

    if os.path.exists(converted_pdf):
        st.success("File converted successfully!")
    else:
        st.error(f"Conversion failed: {converted_pdf}")

# File Uploader for JD (CSV) Uploads
st.header("Job Description(CSV File)Uploader")
jd_uploaded_file = st.file_uploader("Upload JD file", type=["csv"], key="jd")
if jd_uploaded_file:
    jd_path = os.path.join(JD_UPLOAD_FOLDER, jd_uploaded_file.name)
    
    with open(jd_path, "wb") as f:
        f.write(jd_uploaded_file.getbuffer())
    
    st.success("JD file uploaded successfully!")




class FileCleaner:
    def __init__(self, folder_path, csv_file, output_folder):
        """
        Initializes the FileCleaner class with folder paths and CSV file.
        
        Args:
        - folder_path (str): Path to the folder containing PDF files.
        - csv_file (str): Path to the CSV file containing unwanted words.
        - output_folder (str): Path to the folder where cleaned results will be saved.
        """
        self.folder_path = folder_path
        self.csv_file = csv_file
        self.output_folder = output_folder

        # Load unwanted words from the CSV file
        self.unwanted_words = self.load_unwanted_words(self.csv_file)

        # Initialize an empty list to store processed data
        self.data = []

    def load_unwanted_words(self, csv_file):
        """
        Loads unwanted words from a CSV file into a list.
        
        Args:
        - csv_file (str): Path to the CSV file containing unwanted words.
        
        Returns:
        - list: List of unwanted words.
        """
        try:
            df = pd.read_csv(csv_file)
            unwanted_words = df['Unwanted Word'].dropna().tolist()  # Extract unwanted words from the CSV
            return unwanted_words
        except Exception as e:
            print(f"Error loading unwanted words from CSV: {e}")
            return []

    def clean_name_from_filename(self, filename):
        """
        Cleans the file name by removing unwanted words and additional unwanted patterns.
        
        Args:
        - filename (str): The original filename to clean.
        
        Returns:
        - str: Cleaned filename without unwanted words, numbers, special characters, etc.
        """
        name_with_extension = os.path.splitext(filename)[0]
        
        # Create the regex pattern from the unwanted words
        unwanted_pattern = r"(" + "|".join(map(re.escape, self.unwanted_words)) + r")"  # Join words using OR operator
        
        # Additional cleaning patterns
        additional_patterns = [
            r"\d+",  # Remove numbers
            r"[^\w\s]",  # Remove special characters
            r"\bym\b",  # Remove 'ym' as a separate word
        ]
        
        # Apply the cleaning patterns (unwanted words + additional patterns)
        name_with_extension = re.sub(unwanted_pattern, '', name_with_extension, flags=re.IGNORECASE)
        for pattern in additional_patterns:
            name_with_extension = re.sub(pattern, '', name_with_extension)
        
        # Clean extra spaces
        cleaned_name = re.sub(r"\s+", " ", name_with_extension).strip()
        
        return cleaned_name

    def extract_email(self, text):
        """
        Extract email from text using regex.
        
        Args:
        - text (str): The text to extract the email from.
        
        Returns:
        - str: The extracted email, or None if no email is found.
        """
        email_regex = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
        match = re.search(email_regex, text)
        return match.group(0) if match else None

    def extract_phone_number_from_text(self, text):
        """
        Extract phone number from text using regex.
        
        Args:
        - text (str): The text to extract phone numbers from.
        
        Returns:
        - str: The extracted phone number, or None if no phone number is found.
        """
        phone_pattern = r'\+?\d{1,3}[\s\-]?\(?\d{1,4}\)?[\s\-]?\d{1,4}[\s\-]?\d{1,4}[\s\-]?\d{1,9}'
        matches = re.findall(phone_pattern, text)

        for match in matches:
            phone_number = re.sub(r'[^\d+]', '', match)  # Remove non-numeric characters except "+"
            if len(phone_number) >= 10:  # Ensure it's a valid phone number
                return phone_number

        return None

    def extract_text_from_pdf(self, pdf_path):
        """
        Extracts text from a PDF file.
        
        Args:
        - pdf_path (str): Path to the PDF file.
        
        Returns:
        - str: Extracted text from the PDF.
        """
        text = ""
        try:
            with fitz.open(pdf_path) as pdf_reader:
                for page in pdf_reader:
                    text += page.get_text("text") + "\n"
        except Exception as e:
            print(f"Error extracting text from {pdf_path}: {e}")
        
        return text.strip()

    def process_files(self):
        """
        Processes each PDF file in the folder, extracts and cleans relevant data, 
        and stores the results in a list.
        """
        for filename in os.listdir(self.folder_path):
            if filename.lower().endswith('.pdf'):
                file_path = os.path.join(self.folder_path, filename)
                
                # Extract name from the file name
                cleaned_name = self.clean_name_from_filename(filename)
                
                # Extract text from PDF
                text = self.extract_text_from_pdf(file_path)
                
                # Extract email and phone number
                email = self.extract_email(text)
                phone_number = self.extract_phone_number_from_text(text)
                
                # Append the results as a dictionary to the data list
                self.data.append({
                    'Filename': filename,
                    'Name from Filename': cleaned_name,
                    'Email': email,
                    'Phone Number': phone_number
                })

    def save_results(self):
        """
        Converts the processed data into a DataFrame and saves it as a CSV file.
        """
        try:
            # Convert the list of dictionaries into a DataFrame
            df = pd.DataFrame(self.data)

            # Ensure the directory for the output file exists
            os.makedirs(os.path.dirname(self.output_folder), exist_ok=True)
            
            # Save the cleaned results to a CSV file
            df.to_csv(self.output_folder, index=False)

            print(f"File cleaning and CSV generation completed successfully! Results saved to: {self.output_folder}")
        except Exception as e:
            print(f"Error saving results: {e}")

# Example usage
if __name__ == "__main__":
    folder_path = r"C:\Users\inc3061\OneDrive - Texila American University\Documents\Resumepath\Data_Flask_Task\New_Parser\Functions\processed_pdfs"
    csv_file = r"C:\Users\inc3061\OneDrive - Texila American University\Documents\Resumepath\Data_Flask_Task\Resume_Data\scripts\Filename_unwanted.csv"
    output_folder = r"C:\Users\inc3061\OneDrive - Texila American University\Documents\Resumepath\Data_Flask_Task\New_Parser\Functions\test_trainname_demo.csv"

    # Instantiate the FileCleaner class
    file_cleaner = FileCleaner(folder_path, csv_file, output_folder)

    # Process the files and save the results
    file_cleaner.process_files()
    file_cleaner.save_results()


class HeaderExtractor:
    def __init__(self, folder_path, unwanted_words=None, evaluation_warning=None):
        """
        Initializes the HeaderExtractor with the folder path and optional parameters.

        Args:
        - folder_path (str): The path to the folder containing resumes.
        - unwanted_words (list): List of words to be avoided in the first line.
        - evaluation_warning (str): The warning message to avoid in the first line.
        """
        self.folder_path = folder_path
        self.nlp = spacy.load("en_core_web_sm")  # Load spaCy model for NER
        self.unwanted_words = unwanted_words or [
            "CURRICLUM VITAE", "curriculum vitae", "resume", "contact", 
            "personal details", "contact", "Professional Skills", 
            "Name", "SUMMARY", "SKILLS", "EXPERIENCE"
        ]
        self.evaluation_warning = evaluation_warning or "Evaluation Warning: The document was created with Spire.Doc for Python."
        self.data = []

    def extract_text_from_pdf(self, pdf_path):
        """Extracts text from PDF file."""
        return extract_text(pdf_path)

    def extract_text_from_docx(self, docx_path):
        """Extracts text from DOCX file."""
        doc = Document(docx_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text

    def clean_and_get_valid_line(self, text):
        """Cleans the text and returns the first valid line."""
        lines = text.strip().split("\n")
        for line in lines:
            line = line.strip()
            if line and not any(word.lower() in line.lower() for word in self.unwanted_words) and line != self.evaluation_warning:
                return line
        return "No valid line found."

    def extract_name_using_spacy(self, text):
        """Extracts name using spaCy's Named Entity Recognition (NER)."""
        doc = self.nlp(text)
        names = [ent.text for ent in doc.ents if ent.label_ == "PERSON"]
        if names:
            return names[0]
        return None

    def extract_name_using_regex(self, first_line):
        """Extracts name using regex from the first valid line."""
        name_pattern = r'\b([A-Z][a-z]+(?: [A-Z]\.)? [A-Z][a-z]+|[A-Z][a-z]+(?: [A-Z][a-z]+)?)\b'
        matches = re.findall(name_pattern, first_line)
        if matches:
            return matches[0]
        return None

    def extract_full_name(self, resume_path):
        """Extracts full name from the resume, either PDF or DOCX."""
        if resume_path.lower().endswith('.pdf'):
            text = self.extract_text_from_pdf(resume_path)
        elif resume_path.lower().endswith('.docx'):
            text = self.extract_text_from_docx(resume_path)
        else:
            raise ValueError("Unsupported file type. Only PDF and DOCX are supported.")
        
        valid_line = self.clean_and_get_valid_line(text)
        full_name = self.extract_name_using_spacy(valid_line)
        if not full_name:
            full_name = self.extract_name_using_regex(valid_line)

        return full_name, valid_line

    def process_folder_and_store_in_dataframe(self):
        """Processes all files in the folder and stores results in a DataFrame."""
        for root, dirs, files in os.walk(self.folder_path):
            for file in files:
                if file.lower().endswith(('.pdf', '.docx')):
                    file_path = os.path.join(root, file)
                    print(f"Processing file: {file_path}")
                    full_name, extracted_line = self.extract_full_name(file_path)
                    if full_name:
                        print(f"Candidate's Full Name: {full_name}")
                    else:
                        full_name = "Full name not found."
                    
                    self.data.append({"Filename": file, "Full_Name": full_name, "Extracted_Line": extracted_line})

        df = pd.DataFrame(self.data)
        return df

    def convert_to_proper_case(self, text):
        """Converts text to title case."""
        text = str(text)  # Convert to string if it's a number or NaN
        return text.title()

    def save_results_to_csv(self, output_path):
        """Saves the results DataFrame to a CSV file."""
        # Check if the directory exists, if not, create it
        output_dir = os.path.dirname(output_path)
        if not os.path.exists(output_dir):
            print(f"Directory {output_dir} does not exist. Creating...")
            os.makedirs(output_dir)

        # Ensure DataFrame has correct columns
        df = pd.DataFrame(self.data) if self.data else pd.DataFrame(columns=["Filename", "Full_Name", "Extracted_Line"])

        for col in ["Filename", "Full_Name", "Extracted_Line"]:
            if col not in df.columns:
                df[col] = None

        df["Extracted_Line"] = df["Extracted_Line"].fillna("").apply(self.convert_to_proper_case)

        df.to_csv(output_path, index=False)
        print(f"Results saved to: {output_path}")

# Example usage
folder_path = r"C:\Users\inc3061\OneDrive - Texila American University\Documents\Resumepath\Data_Flask_Task\New_Parser\Functions\processed_pdfs"  # Replace with your folder path
output_path = r"C:\Users\inc3061\OneDrive - Texila American University\Documents\Resumepath\Data_Flask_Task\New_Parser\Functions\parsed_headername.csv"  # Path to save the CSV

# Instantiate the HeaderExtractor class
header_extractor = HeaderExtractor(folder_path)

# Process the folder and store results
df = header_extractor.process_folder_and_store_in_dataframe()

# Save the results to a CSV file
header_extractor.save_results_to_csv(output_path)


class ResumeCleaner:
    def __init__(self, unwanted_words_csv, input_csv, output_csv):
        self.unwanted_words_csv = unwanted_words_csv
        self.input_csv = input_csv
        self.output_csv = self.sanitize_filepath(output_csv)
        self.unwanted_words = self.read_unwanted_words()
        
        # Check if input CSV exists and is readable
        if os.path.exists(self.input_csv):
            try:
                self.df = pd.read_csv(self.input_csv)
                if self.df.empty:
                    print(f"Warning: The input CSV file '{self.input_csv}' is empty.")
                    self.df = pd.DataFrame()
            except pd.errors.EmptyDataError:
                print(f"Error: The input CSV file '{self.input_csv}' is empty or unreadable.")
                self.df = pd.DataFrame()
        else:
            print(f"Error: The input CSV file '{self.input_csv}' does not exist.")
            self.df = pd.DataFrame()

    def sanitize_filepath(self, filepath):
        """Sanitizes the filename but retains the full directory path."""
        dir_path, filename = os.path.split(filepath)
        filename = re.sub(r'[<>:"/\\|?*]', '', filename).strip()  # Remove invalid filename characters
        return os.path.join(dir_path, filename) if filename else None
    
    def read_unwanted_words(self):
        """Reads unwanted words from the CSV file."""
        if os.path.exists(self.unwanted_words_csv):
            try:
                unwanted_df = pd.read_csv(self.unwanted_words_csv)
                return unwanted_df['word'].dropna().tolist() if 'word' in unwanted_df.columns else []
            except pd.errors.EmptyDataError:
                print(f"Warning: The unwanted words CSV '{self.unwanted_words_csv}' is empty or unreadable.")
        else:
            print(f"Error: The unwanted words CSV '{self.unwanted_words_csv}' does not exist.")
        return []
    
    def clean_extracted_line(self, extracted_line):
        """Cleans an extracted line by removing unwanted elements."""
        if not isinstance(extracted_line, str):
            return ""
        extracted_line = re.sub(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', '', extracted_line)  # Remove emails
        extracted_line = re.sub(r'(\+?\(?\d{1,4}\)?[\s\-]?)?(\(?\d{1,4}\)?[\s\-]?\d{1,3}[\s\-]?\d{3}[\s\-]?\d{4}|\d{10})', '', extracted_line)  # Remove phone numbers
        extracted_line = re.sub(r'\(.*?\)', '', extracted_line)  # Remove text inside parentheses
        for word in self.unwanted_words:
            extracted_line = re.sub(r'\b' + re.escape(word) + r'\b', '', extracted_line, flags=re.IGNORECASE)
        extracted_line = re.sub(r'[^a-zA-Z\s]', '', extracted_line)  # Keep only alphabetic words
        return extracted_line.strip()
    
    def clean_resumes(self):
        """Cleans the 'Extracted_Line' column in the dataset."""
        if not self.df.empty and 'Extracted_Line' in self.df.columns:
            self.df['Cleaned_Extracted_Line'] = self.df['Extracted_Line'].astype(str).apply(self.clean_extracted_line)
        else:
            print("Error: No valid 'Extracted_Line' column found or DataFrame is empty.")
    
    def save_cleaned_resumes(self):
        """Saves the cleaned data to a new CSV file."""
        if self.df.empty:
            print("Error: No data to save to CSV.")
            return
        
        output_dir = os.path.dirname(self.output_csv)
        os.makedirs(output_dir, exist_ok=True)  # Ensure the directory exists
        try:
            self.df.to_csv(self.output_csv, index=False)
            print(f"Success: Cleaned resumes saved to {self.output_csv}")
        except OSError as e:
            print(f"OS Error: {e}")

# Example usage:
unwanted_words_csv = r"C:\Users\inc3061\OneDrive - Texila American University\Documents\Resumepath\Data_Flask_Task\Resume_Data\scripts\Header_unwanted_new.csv"
input_csv = r"C:\Users\inc3061\OneDrive - Texila American University\Documents\Resumepath\Data_Flask_Task\New_Parser\Functions\parsed_headername.csv"
output_csv = r"C:\Users\inc3061\OneDrive - Texila American University\Documents\Resumepath\Data_Flask_Task\New_Parser\Functions\cleaned_parsed_headername.csv"

resume_cleaner = ResumeCleaner(unwanted_words_csv, input_csv, output_csv)
resume_cleaner.clean_resumes()
resume_cleaner.save_cleaned_resumes()



class NameComparator:
    def __init__(self, filename_csv, headername_csv, output_csv):
        """
        Initializes the NameComparator with paths for filename CSV, headername CSV, and output CSV.
        """
        self.filename_csv = os.path.abspath(filename_csv)
        self.headername_csv = os.path.abspath(headername_csv)
        self.output_csv = self.sanitize_filename(output_csv)

        # Load CSVs with error handling
        self.filename_df = self.load_csv(self.filename_csv, "Filename CSV")
        self.headername_df = self.load_csv(self.headername_csv, "Headername CSV")

    def sanitize_filename(self, path):
        """Removes invalid characters from the filename and ensures the directory exists."""
        directory, filename = os.path.split(path)
        filename = "".join(c for c in filename if c not in '<>:"/\\|?*')  # Remove invalid characters
        os.makedirs(directory, exist_ok=True)  # Ensure directory exists
        return os.path.join(directory, filename)

    def load_csv(self, path, name):
        """Loads a CSV file safely, returning an empty DataFrame if it fails."""
        if not os.path.exists(path):
            print(f"Error: {name} '{path}' does not exist.")
            return pd.DataFrame()
        try:
            df = pd.read_csv(path)
            if df.empty:
                print(f"Warning: {name} '{path}' is empty.")
            return df
        except Exception as e:
            print(f"Error loading {name} '{path}': {e}")
            return pd.DataFrame()

    def clean_extracted_line(self):
        """Cleans the 'Cleaned_Extracted_Line' column."""
        if not self.headername_df.empty:
            self.headername_df['Cleaned_Extracted_Line'] = self.headername_df['Cleaned_Extracted_Line'].apply(
                lambda x: '' if len(str(x)) >= 3 and len(str(x)) < 10 else x
            )

    def merge_dataframes(self):
        """Merges the two DataFrames on the 'Filename' column."""
        if not self.filename_df.empty and not self.headername_df.empty:
            combined_df = pd.merge(self.filename_df, self.headername_df, on='Filename', how='inner')
            return combined_df[['Filename', 'Name from Filename', 'Cleaned_Extracted_Line', 'Email', 'Phone Number']]
        else:
            print("Error: One or both DataFrames are empty. Cannot merge.")
            return pd.DataFrame()

    def check_match(self, row):
        """Compares 'Name from Filename' and 'Cleaned_Extracted_Line'."""
        name_from_filename = str(row['Name from Filename']).lower() if pd.notna(row['Name from Filename']) else ''
        cleaned_line = str(row['Cleaned_Extracted_Line']).lower() if pd.notna(row['Cleaned_Extracted_Line']) else ''
        
        if pd.isna(row['Name from Filename']) and pd.isna(row['Cleaned_Extracted_Line']):
            return ''
        if pd.isna(row['Name from Filename']):
            return row['Cleaned_Extracted_Line']
        if pd.isna(row['Cleaned_Extracted_Line']):
            return row['Name from Filename']
        if cleaned_line in name_from_filename:
            return row['Cleaned_Extracted_Line']
        return row['Name from Filename']

    def compare_names(self):
        """Applies the name comparison logic."""
        merged_df = self.merge_dataframes()
        if not merged_df.empty:
            merged_df['Result'] = merged_df.apply(self.check_match, axis=1)
        else:
            print("Error: No data available to compare.")
        return merged_df

    def save_comparison_results(self, final_df):
        """Saves the comparison results to a CSV file."""
        if final_df.empty:
            print("Error: No comparison results to save.")
            return
        
        try:
            final_df.to_csv(self.output_csv, index=False)
            print(f"Comparison results saved to '{self.output_csv}'")
        except Exception as e:
            print(f"Error saving CSV '{self.output_csv}': {e}")

# Example usage
filename_csv = r"C:\Users\inc3061\OneDrive - Texila American University\Documents\Resumepath\Data_Flask_Task\New_Parser\Functions\test_trainname_demo.csv"
headername_csv = r"C:\Users\inc3061\OneDrive - Texila American University\Documents\Resumepath\Data_Flask_Task\New_Parser\Functions\cleaned_parsed_headername.csv"
output_csv = r"C:\Users\inc3061\OneDrive - Texila American University\Documents\Resumepath\Data_Flask_Task\New_Parser\Functions\Demo_overall_name.csv"

# Run the comparison
name_comparator = NameComparator(filename_csv, headername_csv, output_csv)
name_comparator.clean_extracted_line()
final_df = name_comparator.compare_names()
name_comparator.save_comparison_results(final_df)


class EmailExtractor:
    def __init__(self, folder_path, output_csv_path):
        self.folder_path = folder_path
        self.output_csv_path = self.sanitize_filename(output_csv_path)  # Fix: Sanitize filename

    @staticmethod
    def sanitize_filename(file_path):
        """Sanitizes the filename to remove invalid characters and ensures the directory exists."""
        if not file_path:
            raise ValueError("Error: Output CSV file path is empty or None!")

        directory = os.path.dirname(file_path)
        filename = os.path.basename(file_path)

        # Remove invalid characters
        filename = re.sub(r'[<>:"/\\|?*]', '', filename)

        # Ensure the directory exists
        if directory and not os.path.exists(directory):
            os.makedirs(directory, exist_ok=True)

        # Reconstruct the path
        sanitized_path = os.path.join(directory, filename)

        if not filename:  # If filename is empty after sanitization
            raise ValueError("Error: Sanitized filename is empty!")

        return sanitized_path

    @staticmethod
    def extract_first_email(text):
        """Extracts the first email from the given text."""
        email_pattern = r"([a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+)"
        emails = re.findall(email_pattern, text)
        return emails[0] if emails else None

    @staticmethod
    def extract_text_from_pdf(pdf_file_path):
        """Extracts text from a PDF file."""
        text = ""
        try:
            reader = PdfReader(pdf_file_path)
            for page in reader.pages:
                extracted_text = page.extract_text()
                if extracted_text:
                    text += extracted_text
        except Exception as e:
            print(f"Error extracting text from {pdf_file_path}: {e}")
        return text

    def extract_first_email_from_folder(self):
        """Extracts the first email from each PDF file in the folder."""
        result = []
        total_files = 0
        processed_files = 0

        for filename in os.listdir(self.folder_path):
            total_files += 1
            if filename.lower().endswith('.pdf'):
                pdf_file_path = os.path.join(self.folder_path, filename)
                text = self.extract_text_from_pdf(pdf_file_path)
                first_email = self.extract_first_email(text)

                if first_email:
                    result.append({'Filename': filename, 'first_email': first_email})
                processed_files += 1

        df = pd.DataFrame(result)

        # Logging for debugging
        print(f"Total files: {total_files}")
        print(f"Processed files with emails: {processed_files}")

        return df

    def save_emails_to_csv(self, emails_df):
        """Saves the extracted emails to a CSV file."""
        if emails_df.empty:
            print("No emails found. Skipping file write.")
            return

        if not self.output_csv_path or not os.path.basename(self.output_csv_path):
            print("Error: Output file path is invalid.")
            return

        try:
            print(f"Saving to: {self.output_csv_path}")
            emails_df.to_csv(self.output_csv_path, index=False)
            print(f"Results saved to: {self.output_csv_path}")
        except OSError as e:
            print(f"OS Error while saving CSV: {e}")

# Corrected Output Path Construction
folder_path = r"C:\Users\inc3061\OneDrive - Texila American University\Documents\Resumepath\Data_Flask_Task\New_Parser\Functions\processed_pdfs"

output_csv_path = r"C:\Users\inc3061\OneDrive - Texila American University\Documents\Resumepath\Data_Flask_Task\New_Parser\Functions\email.csv"

# Instantiate and Process
email_extractor = EmailExtractor(folder_path, output_csv_path)
emails_df = email_extractor.extract_first_email_from_folder()
print(emails_df)
email_extractor.save_emails_to_csv(emails_df)


class PhoneNumberExtractor:
    def __init__(self, folder_path, output_csv_path):
        """
        Initializes the PhoneNumberExtractor with folder path and output CSV file path.

        Args:
        - folder_path (str): Path to the folder containing PDF files.
        - output_csv_path (str): Path to save the extracted phone numbers in CSV format.
        """
        self.folder_path = os.path.abspath(folder_path)
        self.output_csv_path = os.path.abspath(output_csv_path)

    @staticmethod
    def extract_phone_number_from_text(text):
        """
        Extracts phone numbers from a given text using regex.

        Args:
        - text (str): Text from which to extract the phone numbers.

        Returns:
        - list: A list of valid phone numbers found in the text.
        """
        phone_pattern = r'\+?\d{1,4}[\s\-]?\(?\d{1,4}\)?[\s\-]?\d{1,3}[\s\-]?\d{3}[\s\-]?\d{4}|\d{10}'
        phone_matches = re.findall(phone_pattern, text)

        phone_numbers = []
        for match in phone_matches:
            phone_number = re.sub(r'\D', '', match)  # Remove non-numeric characters
            if len(phone_number) >= 10:  # Valid phone numbers should have at least 10 digits
                phone_numbers.append(phone_number)

        return phone_numbers if phone_numbers else None

    @staticmethod
    def extract_text_from_pdf(pdf_file_path):
        """
        Extracts text from a PDF file.

        Args:
        - pdf_file_path (str): Path to the PDF file.

        Returns:
        - str: Extracted text from the PDF.
        """
        text = ""
        try:
            with open(pdf_file_path, "rb") as file:
                reader = PdfReader(file)
                for page in reader.pages:
                    extracted_text = page.extract_text()
                    if extracted_text:
                        text += extracted_text + "\n"
        except Exception as e:
            print(f"Error extracting text from {pdf_file_path}: {e}")
        return text.strip()

    def extract_phone_numbers_from_folder(self):
        """
        Extracts phone numbers from all PDFs in the specified folder.

        Returns:
        - pd.DataFrame: A DataFrame with filenames and corresponding phone numbers.
        """
        result = []
        total_files = 0
        processed_files = 0

        if not os.path.exists(self.folder_path):
            print(f"Error: Folder '{self.folder_path}' does not exist.")
            return pd.DataFrame()

        for filename in os.listdir(self.folder_path):
            total_files += 1
            if filename.lower().endswith(".pdf"):
                pdf_file_path = os.path.join(self.folder_path, filename)

                text = self.extract_text_from_pdf(pdf_file_path)
                phone_numbers = self.extract_phone_number_from_text(text)

                if phone_numbers:
                    first_phone_number = phone_numbers[0] if phone_numbers else None
                    result.append({'Filename': filename, 'phone_numbers': phone_numbers, 'first_phone_number': first_phone_number})
                    processed_files += 1

        df = pd.DataFrame(result)

        print(f"Total files: {total_files}")
        print(f"Processed files with phone numbers: {processed_files}")
        return df

    def process_phone_numbers(self, row):
        """
        Processes phone numbers based on the provided logic.

        Args:
        - row (pd.Series): A row from the DataFrame containing phone numbers.

        Returns:
        - str: The processed phone number.
        """
        first_phone_number = row['first_phone_number']
        phone_numbers = row['phone_numbers']

        cleaned_phone_numbers = [re.sub(r'\D', '', num) for num in phone_numbers]

        if not cleaned_phone_numbers:
            return None

        # Custom logic based on phone number length
        if len(first_phone_number) in [4, 6]:  
            return cleaned_phone_numbers[-1][-10:] if len(cleaned_phone_numbers[-1]) >= 10 else None
        elif len(first_phone_number) > 13:
            return cleaned_phone_numbers[-1][-10:] if len(cleaned_phone_numbers[-1]) >= 10 else None
        else:
            return cleaned_phone_numbers[0]

    def save_phone_numbers_to_csv(self, phone_numbers_df):
        """
        Saves the DataFrame containing phone numbers to a CSV file.

        Args:
        - phone_numbers_df (pd.DataFrame): DataFrame containing phone numbers to save.
        """
        if not phone_numbers_df.empty:
            try:
                phone_numbers_df.to_csv(self.output_csv_path, index=False)
                print(f"Results saved to: {self.output_csv_path}")
            except OSError as e:
                print(f"Error saving CSV: {e}")
        else:
            print("No phone numbers to save.")

# Example usage
folder_path = r'C:\Users\inc3061\OneDrive - Texila American University\Documents\Resumepath\Data_Flask_Task\New_Parser\Functions\processed_pdfs'
output_csv_path = r'C:\Users\inc3061\OneDrive - Texila American University\Documents\Resumepath\Data_Flask_Task\New_Parser\Functions\phone_numbers.csv'

# Instantiate the PhoneNumberExtractor class
phone_number_extractor = PhoneNumberExtractor(folder_path, output_csv_path)

# Extract phone numbers from PDFs in the folder
phone_numbers_df = phone_number_extractor.extract_phone_numbers_from_folder()

# Check if the DataFrame is not empty before processing
if not phone_numbers_df.empty:
    phone_numbers_df['processed_phone_number'] = phone_numbers_df.apply(phone_number_extractor.process_phone_numbers, axis=1)
    phone_number_extractor.save_phone_numbers_to_csv(phone_numbers_df)
else:
    print("No phone numbers found in the folder.")



class DataMerger:
    def __init__(self, phone_numbers_csv, email_csv, name_csv, output_csv_path):
        self.phone_numbers_csv = phone_numbers_csv
        self.email_csv = email_csv
        self.name_csv = name_csv
        self.output_csv_path = self.sanitize_filename(output_csv_path)  # Fix: Sanitize filename

        required_columns = {
            "phone_numbers": ["Filename", "phone_numbers", "first_phone_number"],
            "email": ["Filename", "first_email"],
            "name": ["Filename", "Name from Filename", "Cleaned_Extracted_Line", "Result", "Email", "Phone Number"]
        }

        self.phone_numbers_df = self.safe_read_csv(phone_numbers_csv, required_columns["phone_numbers"])
        self.email_df = self.safe_read_csv(email_csv, required_columns["email"])
        self.name_df = self.safe_read_csv(name_csv, required_columns["name"])

    @staticmethod
    def sanitize_filename(file_path):
        """Sanitizes the filename to remove invalid characters."""
        if not file_path:
            raise ValueError("Error: Output CSV file path is empty or None!")

        # Extract directory and filename separately
        directory = os.path.dirname(file_path)
        filename = os.path.basename(file_path)

        # Remove invalid characters from filename
        filename = re.sub(r'[<>:"/\\|?*]', '', filename)

        # Ensure the directory exists
        if directory and not os.path.exists(directory):
            os.makedirs(directory, exist_ok=True)

        # Reconstruct the sanitized file path
        sanitized_path = os.path.join(directory, filename)

        if not filename:  # If filename is empty after sanitization
            raise ValueError("Error: Sanitized filename is empty!")

        return sanitized_path

    def safe_read_csv(self, file_path, required_columns):
        try:
            df = pd.read_csv(file_path)
        except (pd.errors.EmptyDataError, FileNotFoundError):
            df = pd.DataFrame(columns=required_columns)

        for col in required_columns:
            if col not in df.columns:
                df[col] = None  
        
        return df

    def clean_and_prepare_data(self):
        self.phone_numbers_df.rename(columns={'phone_numbers': 'PhoneNumbers_List', 
                                              'first_phone_number': 'Candidates_PhoneNumber'}, inplace=True)
        self.phone_numbers_df = self.phone_numbers_df[['Filename', 'PhoneNumbers_List', 'Candidates_PhoneNumber']]
        self.email_df.rename(columns={'first_email': 'Candidates_Email'}, inplace=True)
        self.name_df.rename(columns={'Name from Filename': 'Name_from_File',
                                     'Cleaned_Extracted_Line': 'Header_Name', 
                                     'Result': 'Candidates_Name', 
                                     'Email': 'Header_Emails', 
                                     'Phone Number': 'Header_Phone'}, inplace=True)
        self.name_df = self.name_df[['Filename', 'Candidates_Name', 'Name_from_File', 'Header_Name', 'Header_Emails', 'Header_Phone']]

    def merge_dataframes(self):
        merged_df = self.name_df.merge(self.email_df, on='Filename', how='left') \
                                .merge(self.phone_numbers_df, on='Filename', how='left')

        merged_df['Header_Emails'] = merged_df['Header_Emails'].fillna(merged_df['Candidates_Email'])
        merged_df['Candidates_PhoneNumber'] = merged_df['Candidates_PhoneNumber'].fillna(merged_df['Header_Phone'])
        merged_df.rename(columns={'Header_Emails': 'Candidates_Emails'}, inplace=True)
        merged_df = merged_df[['Filename', 'Candidates_Name', 'Name_from_File', 'Header_Name', 'Candidates_Emails', 'Candidates_PhoneNumber', 'PhoneNumbers_List']]

        return merged_df

    def save_to_csv(self, merged_df):
        """Saves the merged DataFrame to a CSV file after validation."""
        if merged_df.empty:
            print("No data to save. Skipping file write.")
            return

        if not self.output_csv_path or not os.path.basename(self.output_csv_path):
            print("Error: Output file path is invalid.")
            return

        try:
            print(f"Saving to: {self.output_csv_path}")  
            merged_df.to_csv(self.output_csv_path, index=False)
            print(f"Combined results saved to: {self.output_csv_path}")
        except OSError as e:
            print(f"OS Error while saving CSV: {e}")

    def process_data(self):
        self.clean_and_prepare_data()
        merged_df = self.merge_dataframes()
        self.save_to_csv(merged_df)


# Example usage
phone_numbers_csv = os.path.join(r'C:\Users\inc3061\OneDrive - Texila American University\Documents\Resumepath\Data_Flask_Task\New_Parser\Functions', 'phone_numbers.csv')
email_csv = os.path.join(r'C:\Users\inc3061\OneDrive - Texila American University\Documents\Resumepath\Data_Flask_Task\New_Parser\Functions', 'email.csv')
name_csv = os.path.join(r'C:\Users\inc3061\OneDrive - Texila American University\Documents\Resumepath\Data_Flask_Task\New_Parser\Functions', 'Demo_overall_name.csv')
output_csv_path = os.path.join(r'C:\Users\inc3061\OneDrive - Texila American University\Documents\Resumepath\Data_Flask_Task\New_Parser\Functions', 'overall_Details.csv')

data_merger = DataMerger(phone_numbers_csv, email_csv, name_csv, output_csv_path)
data_merger.process_data()
 



class ExperienceExtractor:
    def __init__(self, folder_path, output_csv_path):
        """
        Initialize the ExperienceExtractor class.

        Args:
            folder_path (str): Path to the folder containing the PDF files.
            output_csv_path (str): Path to save the output CSV file containing experience data.
        """
        self.folder_path = os.path.abspath(folder_path)
        self.output_csv_path = self.sanitize_filename(output_csv_path)
        self.experience_data = []

    def sanitize_filename(self, path):
        """Sanitizes the file path to avoid invalid characters."""
        directory, filename = os.path.split(path)
        filename = re.sub(r'[<>:"/\\|?*]', '', filename)  # Remove invalid characters
        sanitized_path = os.path.join(directory, filename)
        return sanitized_path

    def extract_experience_from_filename(self, filename):
        """
        Extract total experience from a filename using regex patterns.

        Args:
            filename (str): The name of the file from which to extract experience.

        Returns:
            tuple: A tuple of (years, months) representing total experience.
        """
        total_experience = None

        # Pattern 1: Numbers before "year", "years", "yr", "yrs" or "m"
        pattern_1 = r'(\d+(\.\d+)?)\s*(year|years|yr|yrs|m)'
        match_1 = re.search(pattern_1, filename, re.IGNORECASE)

        if match_1:
            years = int(float(match_1.group(1)))  # Convert the value to int years
            months = int((float(match_1.group(1)) - years) * 12)  # Convert decimal to months
            total_experience = (years, months)

        # Pattern 2: Numbers in square brackets like [10y_8m]
        pattern_2 = r'\[(\d+)(y|yrs|yr?)?[_-](\d+)(m)?\]'
        match_2 = re.search(pattern_2, filename)

        if match_2:
            years_in_bracket = int(match_2.group(1))
            months_in_bracket = int(match_2.group(3))
            total_experience = (years_in_bracket, months_in_bracket)

        return total_experience

    def process_files_in_folder(self):
        """
        Iterate through files in the specified folder and extract experience information.

        Returns:
            list: A list of lists containing filenames and their corresponding experience.
        """
        if not os.path.exists(self.folder_path):
            print(f"Error: Folder '{self.folder_path}' does not exist.")
            return

        for filename in os.listdir(self.folder_path):
            if filename.endswith(".pdf"):
                experience = self.extract_experience_from_filename(filename)
                if experience:
                    self.experience_data.append([filename, f"{experience[0]} years, {experience[1]} months"])

    def save_experience_to_csv(self):
        """
        Save the extracted experience data to a CSV file.
        """
        if not self.experience_data:
            print("No experience data to save.")
            return

        output_dir = os.path.dirname(self.output_csv_path)
        os.makedirs(output_dir, exist_ok=True)  # Ensure output directory exists

        try:
            df = pd.DataFrame(self.experience_data, columns=["Filename", "Total_Experience"])
            df.to_csv(self.output_csv_path, index=False)
            print(f"Experience data saved to: {self.output_csv_path}")
        except OSError as e:
            print(f"Error saving CSV: {e}")

    def extract_and_save(self):
        """
        Extract experience data from files in the folder and save it to CSV.
        """
        self.process_files_in_folder()
        self.save_experience_to_csv()

# Example usage
folder_path = r'C:\Users\inc3061\OneDrive - Texila American University\Documents\Resumepath\Data_Flask_Task\New_Parser\Functions\processed_pdfs'
output_csv_path = r'C:\Users\inc3061\OneDrive - Texila American University\Documents\Resumepath\Data_Flask_Task\New_Parser\Functions\Fileexp_data.csv'

experience_extractor = ExperienceExtractor(folder_path, output_csv_path)
experience_extractor.extract_and_save()


class RandomExperienceExtractor:
    def __init__(self, folder_path, output_csv_path):
        self.folder_path = folder_path
        self.output_csv_path = output_csv_path

    def extract_experience_from_pdf(self, file_path):
        try:
            reader = PdfReader(file_path)
            if len(reader.pages) == 0:
                return None
            
            first_page_text = reader.pages[0].extract_text()
            if not first_page_text:
                return None

            experience_pattern = r"\b(?:over|more than|at least|greater than)?\s*(\d+(\.\d+)?)\s*(?:years?|Yrs?)\s+of(?:\s+[a-zA-Z]+)*\s*(experience|progressively\s+in\s+[a-zA-Z\s]+|professional\s+experience|rich\s+and\s+extensive\s+experience)\b"
            matches = re.findall(experience_pattern, first_page_text, re.IGNORECASE)

            additional_experience_pattern = r"\b(?:over|more than|at least|greater than)?\s*(\d+(\.\d+)?)\s*(?:years?|Yrs?)\s+(?:of\s+[a-zA-Z]+)*\s+professional\s+practice\b"
            additional_matches = re.findall(additional_experience_pattern, first_page_text, re.IGNORECASE)

            total_matches = matches + additional_matches

            return list(set([match[0] for match in total_matches])) if total_matches else ["No experience found"]
        except Exception as e:
            print(f"Error processing {file_path}: {e}")
            return ["No experience found"]

    def process_folder(self):
        data = []
        for filename in os.listdir(self.folder_path):
            if filename.endswith('.pdf'):
                file_path = os.path.join(self.folder_path, filename)
                experiences = self.extract_experience_from_pdf(file_path)
                
                experience_str = ', '.join(experiences) if experiences else "No experience found"
                data.append({"Filename": filename, "experience_list": experience_str})
        
        df = pd.DataFrame(data)
        
        if df.empty:
            print("Warning: No experience data extracted. The DataFrame is empty.")
        else:
            print("Processed DataFrame:\n", df.head()) 
        
        return df

    def filter_and_convert_experience(self, df):
        if "experience_list" not in df.columns:
            print("Error: 'experience_list' column is missing from the DataFrame.")
            return pd.DataFrame()  # Return an empty DataFrame to prevent further errors
        
        filtered_df = df[df['experience_list'] != 'No experience found']

        def convert_experience_to_year_month(exp_str):
            try:
                exp_num = float(exp_str)
                years = int(exp_num)
                months = round((exp_num - years) * 12)
                return f"{years} years, {months} months"
            except ValueError:
                pass  
            
            year_month_pattern = r'(\d+)(?:\s?years?|yr)?(?:\s?(\d+)\s?months?)?'
            match = re.match(year_month_pattern, str(exp_str).lower())

            if match:
                years = int(match.group(1))
                months = int(match.group(2)) if match.group(2) else 0
                return f"{years} years, {months} months"

            try:
                years = int(exp_str)
                return f"{years} years, 0 months"
            except ValueError:
                return None  

        if not filtered_df.empty:
            filtered_df['formatted_experience'] = filtered_df['experience_list'].apply(convert_experience_to_year_month)
        
        return filtered_df

    def save_experience_to_csv(self, filtered_df):
        if filtered_df.empty:
            print("Warning: No data to save. Skipping CSV generation.")
            return
        filtered_df.to_csv(self.output_csv_path, index=False)
        print(f"Experience data saved to {self.output_csv_path}")

    def extract_and_process_experience(self):
        df = self.process_folder()
        if df.empty:
            print("Error: No data extracted. Exiting process.")
            return
        filtered_df = self.filter_and_convert_experience(df)
        self.save_experience_to_csv(filtered_df)

# Paths (Change as needed)
folder_path = r'C:\Users\inc3061\OneDrive - Texila American University\Documents\Resumepath\Data_Flask_Task\New_Parser\Functions\processed_pdfs'
output_csv_path = r'C:\Users\inc3061\OneDrive - Texila American University\Documents\Resumepath\Data_Flask_Task\New_Parser\Functions\Randomexp_data.csv'

random_experience_extractor = RandomExperienceExtractor(folder_path, output_csv_path)
random_experience_extractor.extract_and_process_experience()


import os
import re
import pandas as pd
import fitz  # PyMuPDF

class ResumeExperienceExtractor:
    def __init__(self, folder_path, search_words_file, output_csv_path):
        """
        Initialize the class with the necessary parameters.

        Args:
        - folder_path (str): Path to the folder containing PDF files.
        - search_words_file (str): Path to the CSV file containing search words.
        - output_csv_path (str): Path where the output CSV file will be saved.
        """
        self.folder_path = os.path.abspath(folder_path)
        self.search_words_file = os.path.abspath(search_words_file)
        self.output_csv_path = self.sanitize_filename(output_csv_path)  # Sanitize output file path
        self.search_words = self.load_search_words()

    def sanitize_filename(self, path):
        """Sanitizes the file path to avoid invalid characters."""
        directory, filename = os.path.split(path)
        filename = re.sub(r'[<>:"/\\|?*]', '', filename)  # Remove invalid characters
        sanitized_path = os.path.join(directory, filename)
        return sanitized_path

    def load_search_words(self):
        """Load search words from a CSV file."""
        try:
            search_words_df = pd.read_csv(self.search_words_file)
            if 'Search_words' in search_words_df.columns:
                return search_words_df['Search_words'].dropna().tolist()
            else:
                print("Error: 'Search_words' column missing in CSV.")
                return []
        except Exception as e:
            print(f"Error loading search words: {e}")
            return []

    def extract_experience(self, text):
        """Extract years of experience from the text."""
        experience_patterns = [
            r"\b(?:over|more than|at least|greater than)?\s*(\d+(\.\d+)?)\s*(?:years?|Yrs?)\s+(?:of\s+[a-zA-Z]+)*\s*(?:experience|progressive\s+experience|professional\s+experience|rich\s+and\s+extensive\s+experience)\b",
            r"\b(?:over|more than|at least|greater than)?\s*(\d+(\.\d+)?)\s*(?:years?|Yrs?)\s+(?:of\s+[a-zA-Z]+)*\s+professional\s+practice\b"
        ]

        total_exp = 0.0

        for pattern in experience_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                try:
                    total_exp += float(match[0])
                except ValueError:
                    continue

        for word in self.search_words:
            pattern = rf'(\d+)\s*(?:years?|Yrs?)\s*(?=\s+{re.escape(word)})'
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                try:
                    total_exp += int(match)
                except ValueError:
                    continue

        return round(total_exp, 1) if total_exp > 0 else None

    def process_pdfs(self):
        """Process PDF files in the given folder and extract experience."""
        if not os.path.exists(self.folder_path):
            print(f"Error: Folder path '{self.folder_path}' does not exist.")
            return pd.DataFrame(columns=['Filename', 'total_Exp'])

        data = []

        for filename in os.listdir(self.folder_path):
            if filename.lower().endswith(".pdf"):
                file_path = os.path.join(self.folder_path, filename)
                
                try:
                    doc = fitz.open(file_path)
                    text = "\n".join([page.get_text("text") for page in doc])
                    total_exp = self.extract_experience(text)

                    if total_exp is not None:
                        data.append({'Filename': filename, 'total_Exp': total_exp})
                
                except Exception as e:
                    print(f"Error processing file {filename}: {e}")

        df = pd.DataFrame(data, columns=['Filename', 'total_Exp'])

        if df.empty:
            print("Warning: No valid experience data extracted.")
        
        return df

    def save_to_csv(self, df):
        """Save the extracted experience data to a CSV file."""
        if df.empty:
            print("Warning: No data to save.")
            return

        output_dir = os.path.dirname(self.output_csv_path)
        os.makedirs(output_dir, exist_ok=True)  # Ensure the directory exists

        try:
            df.to_csv(self.output_csv_path, index=False)
            print(f"Results saved to {self.output_csv_path}")
        except Exception as e:
            print(f"Error saving CSV: {e}")

    def run(self):
        """Run the entire process."""
        df = self.process_pdfs()

        if 'total_Exp' in df.columns and not df.empty:
            df = df[df['total_Exp'] > 0.0]
        else:
            print("Warning: 'total_Exp' column not found or DataFrame is empty.")

        self.save_to_csv(df)


# Example Usage
folder_path = r'C:\Users\inc3061\OneDrive - Texila American University\Documents\Resumepath\Data_Flask_Task\New_Parser\Functions\processed_pdfs'
search_words_file = r'C:\Users\inc3061\OneDrive - Texila American University\Documents\Resumepath\Data_Flask_Task\Resume_Data\scripts\CSV_serachExpheader.csv'
output_csv_path = r'C:\Users\inc3061\OneDrive - Texila American University\Documents\Resumepath\Data_Flask_Task\New_Parser\Functions\csv_output_experience.csv'

resume_extractor = ResumeExperienceExtractor(folder_path, search_words_file, output_csv_path)
resume_extractor.run()


class ExperienceProcessor:
    def __init__(self, filename_csv, search_csv, content_csv, output_csv):
        """Initializes the ExperienceProcessor with paths to input/output CSV files."""
        self.filename_df = self.read_and_clean_csv(filename_csv)
        self.search_df = self.read_and_clean_csv(search_csv)
        self.content_df = self.read_and_clean_csv(content_csv)
        self.output_csv_path = self.sanitize_path(output_csv)  # Fix: Ensure valid path

    def sanitize_path(self, file_path):
        """Sanitizes file paths and ensures directory exists."""
        if not file_path:
            raise ValueError("Error: Output CSV file path is empty or None!")

        directory = os.path.dirname(file_path)
        filename = os.path.basename(file_path)

        # Remove invalid characters from filename
        filename = re.sub(r'[<>:"/\\|?*]', '', filename)

        # Ensure directory exists
        if directory and not os.path.exists(directory):
            os.makedirs(directory, exist_ok=True)

        return os.path.join(directory, filename)

    def read_and_clean_csv(self, file_path):
        """Reads a CSV file and strips whitespace from column names."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Error: File not found - {file_path}")

        df = pd.read_csv(file_path)
        df.columns = df.columns.str.strip()  # Clean column names
        return df

    def merge_dataframes(self):
        """Merges filename, search, and content DataFrames on 'Filename'."""
        if 'Filename' not in self.filename_df.columns:
            raise KeyError("Error: 'Filename' column missing in filename CSV.")

        merge_df = pd.merge(self.filename_df, self.search_df, on='Filename', how='outer')
        merged_df = pd.merge(merge_df, self.content_df, on='Filename', how='outer')
        
        print("Merged DataFrame Columns:", merged_df.columns.tolist())  # Debugging
        return merged_df

    def convert_to_decimal_years(self, exp):
        """Converts 'years, months' experience format to decimal years."""
        exp = str(exp)
        year_month_pattern = r"(\d+)\s*years?,?\s*(\d+)?\s*months?|(\d+)\s*years?"

        match = re.search(year_month_pattern, exp)
        if match:
            if match.group(1) and match.group(2):  # Both years and months present
                years = int(match.group(1))
                months = int(match.group(2)) if match.group(2) else 0
                decimal_years = years + (months / 12)
            elif match.group(3):  # Only years present
                decimal_years = int(match.group(3))
            return round(decimal_years, 1)
        return None  # Return None if format is invalid

    def apply_decimal_conversion(self, df, column_name):
        """Applies decimal conversion to a specific column in DataFrame."""
        if column_name in df.columns:
            df[f'{column_name}_decimal'] = df[column_name].apply(self.convert_to_decimal_years)
        else:
            print(f"Warning: Column '{column_name}' not found in DataFrame.")
        return df

    def get_overall_exp(self, row):
        """Assigns 'overall_exp' based on available experience columns."""
        if pd.notna(row['Filenme_exp']):
            return row['Filenme_exp']
        elif pd.notna(row['Content_exp']):
            return row['Content_exp']
        elif pd.notna(row['Header_exp']):
            return row['Header_exp']
        return np.nan  # Default to NaN if no values are available

    def generate_final_dataframe(self):
        """Processes experience data and saves the final output CSV."""
        merged_df = self.merge_dataframes()

        # Ensure the 'total_exp' column exists before applying conversion
        if 'total_exp' in merged_df.columns:
            merged_df = self.apply_decimal_conversion(merged_df, 'total_exp')
        else:
            print("Warning: 'total_exp' column not found in DataFrame.")

        # Rename columns for final format
        column_mapping = {
            'Filename': 'Filename',
            'total_exp': 'Filenme_exp',
            'total_Exp': 'Header_exp',
            'experience_list': 'Content_exp'
        }

        # Select only the necessary columns
        existing_columns = [col for col in column_mapping.keys() if col in merged_df.columns]
        merged_df = merged_df[existing_columns]

        # Rename the columns
        merged_df.rename(columns=column_mapping, inplace=True)

        # Apply function to compute 'overall_exp'
        merged_df['overall_exp'] = merged_df.apply(self.get_overall_exp, axis=1)

        # Save the final DataFrame to CSV
        output_columns = ['Filename', 'overall_exp']
        merged_df = merged_df[output_columns]
        merged_df.to_csv(self.output_csv_path, index=False)
        print(f"Results saved to {self.output_csv_path}")

#Example Usage
def main():
    processor = ExperienceProcessor(
        filename_csv=r'C:\Users\inc3061\OneDrive - Texila American University\Documents\Resumepath\Data_Flask_Task\New_Parser\Functions\experience_data.csv',
        search_csv=r'C:\Users\inc3061\OneDrive - Texila American University\Documents\Resumepath\Data_Flask_Task\New_Parser\Functions\Randomexp_data.csv',
        content_csv=r'C:\Users\inc3061\OneDrive - Texila American University\Documents\Resumepath\Data_Flask_Task\New_Parser\Functions\csv_output_experience.csv',
        output_csv=r'C:\Users\inc3061\OneDrive - Texila American University\Documents\Resumepath\Data_Flask_Task\New_Parser\Functions\Compare_exp.csv'
    )
    
    processor.generate_final_dataframe()

# Run the script
if __name__ == "__main__":
    main()

class PDFMatcher:
    def __init__(self, pdf_folder_path, jd_folder_path, output_csv_path):
        self.pdf_folder_path = os.path.abspath(pdf_folder_path)
        self.jd_folder_path = os.path.abspath(jd_folder_path)
        self.output_csv_path = self.sanitize_filename(output_csv_path)
        self.results = []
        self.jd_csv_path = self.get_latest_csv_from_folder()

    def sanitize_filename(self, path):
        """Removes invalid characters from filenames."""
        directory, filename = os.path.split(path)
        filename = "".join(c for c in filename if c not in '<>:"/\\|?*')
        return os.path.join(directory, filename)

    def get_latest_csv_from_folder(self):
        """Finds the latest CSV file in the given JD folder."""
        if not os.path.exists(self.jd_folder_path):
            print(f"Error: JD folder '{self.jd_folder_path}' does not exist.")
            return None
        
        csv_files = [f for f in os.listdir(self.jd_folder_path) if f.endswith('.csv')]
        
        if not csv_files:
            print(f"Error: No CSV file found in {self.jd_folder_path}")
            return None

        csv_files.sort(key=lambda f: os.path.getmtime(os.path.join(self.jd_folder_path, f)), reverse=True)
        latest_csv = os.path.join(self.jd_folder_path, csv_files[0])
        
        print(f"Using JD CSV: {latest_csv}")
        return latest_csv

    def extract_text_from_pdf(self, pdf_path):
        """Extracts text from a PDF."""
        text = ""
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    extracted_text = page.extract_text()
                    if extracted_text:
                        text += extracted_text
        except Exception as e:
            print(f"Error reading PDF '{pdf_path}': {e}")
        
        return text.replace("Evaluation Warning: The document was created with Spire.Doc for Python.", "") if text else ""

    def calculate_match_percentage(self, text, search_word):
        """Calculates match percentage using fuzzy string matching."""
        return fuzz.partial_ratio(text.lower(), search_word.lower()) if text and search_word else 0

    def process_pdfs_and_find_matches(self):
        """Processes PDFs and finds matching words."""
        if not self.jd_csv_path or not os.path.exists(self.jd_csv_path):
            print("No valid JD CSV file found. Aborting process.")
            return pd.DataFrame()

        try:
            sql_jd_df = pd.read_csv(self.jd_csv_path)
            if 'Skills' not in sql_jd_df.columns or 'Roles' not in sql_jd_df.columns:
                print("Error: JD CSV missing required columns ('Skills', 'Roles').")
                return pd.DataFrame()
        except Exception as e:
            print(f"Error reading JD CSV: {e}")
            return pd.DataFrame()

        if not os.path.exists(self.pdf_folder_path):
            print(f"Error: PDF folder '{self.pdf_folder_path}' does not exist.")
            return pd.DataFrame()

        for filename in os.listdir(self.pdf_folder_path):
            if filename.lower().endswith('.pdf'):
                pdf_path = os.path.join(self.pdf_folder_path, filename)
                pdf_text = self.extract_text_from_pdf(pdf_path)

                for _, row in sql_jd_df.iterrows():
                    search_word = str(row.get('Skills', '')).strip()
                    role = str(row.get('Roles', 'Unknown')).strip()

                    if search_word:
                        match_percentage = self.calculate_match_percentage(pdf_text, search_word)

                        if match_percentage > 50:  # Match threshold
                            self.results.append({
                                'Filename': filename,
                                'Match Word': search_word,
                                'Match Percentage': match_percentage,
                                'Role': role
                            })

        results_df = pd.DataFrame(self.results)
        if not results_df.empty:
            os.makedirs(os.path.dirname(self.output_csv_path), exist_ok=True)
            results_df.to_csv(self.output_csv_path, index=False)
            print(f"Matches found. Results saved to '{self.output_csv_path}'.")
        else:
            print("No strong matches found.")

        return results_df

# Example usage
def main():
    pdf_folder_path = r'C:\Users\inc3061\OneDrive - Texila American University\Documents\Resumepath\Data_Flask_Task\New_Parser\Functions\processed_pdfs'
    jd_folder_path = r'C:\Users\inc3061\OneDrive - Texila American University\Documents\Resumepath\Data_Flask_Task\New_Parser\Functions\JD_uploads'
    output_csv_path = r'C:\Users\inc3061\OneDrive - Texila American University\Documents\Resumepath\Data_Flask_Task\New_Parser\Functions\JD_Match.csv'

    pdf_matcher = PDFMatcher(pdf_folder_path, jd_folder_path, output_csv_path)
    results_df = pdf_matcher.process_pdfs_and_find_matches()

    if not results_df.empty:
        print(results_df)

if __name__ == "__main__":
    main()

class ResumeJDMatcher:
    def __init__(self, directory, output_csv_path):
        self.directory = os.path.abspath(directory)
        self.details_csv_path = self.get_latest_file("overall_Details")
        self.jd_match_csv_path = self.get_latest_file("JD_Match")
        
        # Validate and fix output path
        if not output_csv_path.strip():
            raise ValueError("Invalid output_csv_path provided.")
        
        self.output_csv_path = os.path.abspath(output_csv_path)
        self.merged_df = None

    def get_latest_file(self, keyword):
        """Finds the most recent file containing the keyword in the directory."""
        files = glob(os.path.join(self.directory, f"*{keyword}*.csv"))
        if not files:
            print(f"No files found for {keyword}. Returning None.")
            return None  # Ensure None is handled correctly
    
        latest_file = max(files, key=os.path.getmtime)  # Get the most recent file
        print(f"Latest {keyword} file: {latest_file}")
    
        # Validate path correctness
        if os.path.exists(latest_file):
            return latest_file
        else:
            print(f"Warning: File {latest_file} does not exist.")
            return None

    def load_data(self):
        """Loads the CSV files into pandas DataFrames."""
        self.df1 = self.load_csv(self.details_csv_path, "overall_Details")
        self.df2 = self.load_csv(self.jd_match_csv_path, "JD_Match")

    def load_csv(self, file_path, label):
        """Helper function to load a CSV file safely."""
        if file_path and os.path.exists(file_path):
            try:
                df = pd.read_csv(file_path, encoding='utf-8')
                if df.empty:
                    print(f"Warning: {label} CSV is empty.")
                return df
            except Exception as e:
                print(f"Error loading {label} CSV: {e}")
        else:
            print(f"Error: {label} CSV not found.")
        return None

    def merge_data(self):
        """Merges the two DataFrames on the 'Filename' column."""
        if self.df1 is None or self.df2 is None:
            print("Error: One or both dataframes are not loaded properly. Cannot merge.")
            return None

        # Ensure 'Filename' column exists
        if "Filename" not in self.df1.columns or "Filename" not in self.df2.columns:
            print("Error: 'Filename' column missing in one of the DataFrames.")
            return None

        # Clean and standardize filenames
        self.df1["Filename"] = self.df1["Filename"].astype(str).str.strip().str.lower()
        self.df2["Filename"] = self.df2["Filename"].astype(str).str.strip().str.lower()

        # Merge DataFrames
        self.merged_df = pd.merge(self.df1, self.df2, on="Filename", how="left")
        print("DataFrames merged successfully. Sample output:")
        print(self.merged_df.head(2))
        return self.merged_df

    def save_output(self):
        """Saves the merged DataFrame to a CSV file."""
        if self.merged_df is not None and not self.merged_df.empty:
            output_dir = os.path.dirname(self.output_csv_path)
    
            if not output_dir or not os.path.exists(output_dir):
                print(f"Creating output directory: {output_dir}")
                os.makedirs(output_dir, exist_ok=True)
    
            try:
                self.merged_df.to_csv(self.output_csv_path, index=False, encoding='utf-8')
                print(f"Merged data saved to '{self.output_csv_path}'.")
            except OSError as e:
                print(f"Error saving CSV: {e}")
        else:
            print("Error: No merged data to save.")

# Example usage
def main():
    directory = r'C:\\Users\\inc3061\\OneDrive - Texila American University\\Documents\\Resumepath\\Data_Flask_Task\\New_Parser\\Functions'
    output_csv_path = os.path.join(directory, 'merged_output.csv')

    matcher = ResumeJDMatcher(directory, output_csv_path)
    matcher.load_data()
    merged_df = matcher.merge_data()

    if merged_df is not None and not merged_df.empty:
        print(merged_df.head())  # Print first few rows for verification
        matcher.save_output()
    else:
        print("Merged DataFrame is empty or None.")

if __name__ == "__main__":
    main()


# Cleanup Functions
def cleanup_non_pdfs():
    for file in os.listdir(TEMP_PDF_FOLDER):
        if not file.lower().endswith(".pdf"):
            os.remove(os.path.join(TEMP_PDF_FOLDER, file))
    st.sidebar.success("Non-PDF files removed!")

def cleanup_all_pdfs():
    for file in os.listdir(TEMP_PDF_FOLDER):
        if file.lower().endswith(".pdf"):
            os.remove(os.path.join(TEMP_PDF_FOLDER, file))
    st.sidebar.success("All PDFs removed!")

def cleanup_all_jd_files():
    """Remove all files from JD Uploads folder."""
    for filename in os.listdir(JD_UPLOAD_FOLDER):
        file_path = os.path.join(JD_UPLOAD_FOLDER, filename)
        if os.path.isfile(file_path):  # Check if it's a file (not a folder)
            os.remove(file_path)

    st.sidebar.success("All JD files removed!")
# Sidebar navigation and buttons
st.sidebar.title("File Management")
st.sidebar.button("Clear Non-PDF Files", on_click=cleanup_non_pdfs)
st.sidebar.button("Clear All PDFs", on_click=cleanup_all_pdfs)
st.sidebar.button("Clear All JD Uploads", on_click=cleanup_all_jd_files)

# Define the path for Merged_out.csv
MERGED_OUT_PATH = r"C:\Users\inc3061\OneDrive - Texila American University\Documents\Resumepath\Data_Flask_Task\New_Parser\Functions\merged_output.csv"

# Check if the file exists
if os.path.exists(MERGED_OUT_PATH):
    # Load the Merged_out CSV file
    merged_df = pd.read_csv(MERGED_OUT_PATH)

    # Display the table
    st.header("Merged Job Description & Resume Data")
    st.dataframe(merged_df)  # Interactive table

    # Convert DataFrame to CSV format for download
    merged_csv = merged_df.to_csv(index=False).encode("utf-8")

    # Buttons for View and Download
    st.download_button(
        label="Download Merged Data",
        data=merged_csv,
        file_name="Merged_out.csv",
        mime="text/csv"
    )
else:
    st.warning("⚠️ Merged_out.csv not found. Please ensure the file is generated.")



